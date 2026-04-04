from io import BytesIO

from PIL import Image
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.main import app


client = TestClient(app)


def test_extract_text_file() -> None:
    response = client.post(
        '/extract',
        files={'file': ('hello.txt', b'hello\n\nworld', 'text/plain')},
        data={'include_chunks': 'true'},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['metadata']['source_type'] == 'text'
    assert 'hello' in payload['raw_text']
    assert len(payload['chunks']) >= 1


def test_extract_blank_text_file_returns_partial() -> None:
    response = client.post(
        '/extract',
        files={'file': ('blank.txt', b'   \n\n', 'text/plain')},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['extraction']['status'] == 'partial'
    warning_codes = {warning['code'] for warning in payload['extraction']['warnings']}
    assert 'text_no_content_detected' in warning_codes
    assert payload['segments'] == []
    assert payload['chunks'] == []
    assert payload['extra']['content_detected'] is False
    assert payload['extra']['empty_content'] is True
    assert payload['extra']['partial_reason'] == 'empty_content'
    assert payload['extra']['warning_codes'] == ['text_no_content_detected']


def test_extract_blank_docx_returns_partial() -> None:
    buffer = BytesIO()
    document = Document()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/extract',
        files={
            'file': (
                'blank.docx',
                buffer,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['metadata']['source_type'] == 'docx'
    assert payload['extraction']['status'] == 'partial'
    warning_codes = {warning['code'] for warning in payload['extraction']['warnings']}
    assert 'docx_no_text_detected' in warning_codes
    assert payload['extra']['content_detected'] is False
    assert payload['extra']['empty_content'] is True
    assert payload['extra']['partial_reason'] == 'empty_content'



def test_extract_docx_headings_and_list_items_preserve_structure_metadata() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_heading('Project Plan', level=1)
    document.add_paragraph('First bullet', style='List Bullet')
    document.add_paragraph('Normal paragraph')
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/extract',
        files={
            'file': (
                'structured.docx',
                buffer,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    section_segments = [segment for segment in payload['segments'] if segment['type'] == 'section']
    paragraph_segments = [segment for segment in payload['segments'] if segment['type'] == 'paragraph']
    assert any(segment['metadata'].get('structure') == 'heading' for segment in section_segments)
    assert any(segment['metadata'].get('heading_level') == 1 for segment in section_segments)
    assert any(segment['metadata'].get('structure') == 'list_item' for segment in paragraph_segments)



def test_extract_docx_with_table_preserves_table_segment() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph('Intro paragraph')
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = 'Role'
    table.cell(1, 0).text = 'Drew'
    table.cell(1, 1).text = 'Admin'
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/extract',
        files={
            'file': (
                'table.docx',
                buffer,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['extraction']['status'] == 'success'
    segment_types = [segment['type'] for segment in payload['segments']]
    assert 'paragraph' in segment_types
    assert 'table' in segment_types
    table_segments = [segment for segment in payload['segments'] if segment['type'] == 'table']
    assert any('Name | Role' in segment['text'] for segment in table_segments)
    assert table_segments[0]['metadata']['table_index'] == 1
    assert payload['extra']['table_segment_count'] >= 1
    assert 'Intro paragraph' in payload['raw_text']
    assert 'Drew | Admin' in payload['raw_text']



def test_extract_blank_xlsx_returns_partial() -> None:
    workbook = Workbook()
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/extract',
        files={
            'file': (
                'blank.xlsx',
                buffer,
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['metadata']['source_type'] == 'xlsx'
    assert payload['extraction']['status'] == 'partial'
    warning_codes = {warning['code'] for warning in payload['extraction']['warnings']}
    assert 'xlsx_no_text_detected' in warning_codes
    assert payload['extra']['content_detected'] is False
    assert payload['extra']['empty_content'] is True
    assert payload['extra']['partial_reason'] == 'empty_content'



def test_extract_xlsx_includes_row_level_table_segments() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'People'
    sheet.append(['Name', 'Role'])
    sheet.append(['Drew', 'Admin'])
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/extract',
        files={
            'file': (
                'people.xlsx',
                buffer,
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload['extraction']['status'] == 'success'
    sheet_segments = [segment for segment in payload['segments'] if segment['type'] == 'sheet']
    table_segments = [segment for segment in payload['segments'] if segment['type'] == 'table']
    assert len(sheet_segments) == 1
    assert len(table_segments) == 2
    assert payload['extra']['non_empty_sheet_count'] == 1
    assert payload['extra']['table_segment_count'] == 2
    assert sheet_segments[0]['metadata']['header_values'] == ['Name', 'Role']
    assert table_segments[0]['metadata']['sheet_name'] == 'People'
    assert table_segments[0]['metadata']['row_index'] == 1
    assert table_segments[0]['metadata']['is_header'] is True
    assert table_segments[1]['metadata']['header_values'] == ['Name', 'Role']
    assert table_segments[1]['metadata']['row_mapping'] == {'Name': 'Drew', 'Role': 'Admin'}
    assert table_segments[1]['text'] == 'Drew | Admin'



def test_extract_image_returns_ocr_metadata(monkeypatch, tmp_path) -> None:
    monkeypatch.setattr('app.extractors.image_ocr.tesseract_available', lambda: True)
    monkeypatch.setattr('app.ocr.pytesseract.image_to_string', lambda image: 'hello image')

    image_path = tmp_path / 'scan.png'
    Image.new('RGB', (40, 40), color='white').save(image_path)

    with image_path.open('rb') as fh:
        response = client.post(
            '/extract',
            files={'file': ('scan.png', fh.read(), 'image/png')},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload['metadata']['source_type'] == 'image'
    assert payload['extra']['ocr_attempted'] is True
    assert payload['extra']['result_source'] == 'ocr'
    assert payload['extra']['processed_mode'] == 'L'
    assert payload['extra']['selected_ocr_pass']
    assert len(payload['extra']['ocr_passes']) >= 3
    assert payload['extra']['non_empty_page_count'] == 1
